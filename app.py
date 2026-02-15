from flask import (
    Flask,
    render_template,
    request,
    redirect,
    url_for,
    session,
    flash,
    send_file,
    jsonify,
)
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
from datetime import datetime
import sqlite3
import os
import json

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


app = Flask(__name__)
app.secret_key = "change-this-secret"

UPLOAD_FOLDER = "static"
# Use the actual template present in the project
TEMPLATE_PATH = "word_templates/college_letterhead.docx"
GENERATED_FOLDER = "generated_reports"

# ---------------- PSO CONSTANTS ----------------
# Programme Specific Outcomes text used in generated reports.
PSO1_TEXT = (
    "PSO1: An ability to apply the theoretical concepts and practical knowledge of "
    "Information Technology in the analysis, design, development, and management of "
    "information processing systems and applications in the interdisciplinary domain "
    "to understand professional, business processes, ethical, legal, security, and "
    "social issues and responsibilities."
)

PSO2_TEXT = (
    "PSO2: An ability to analyze a problem and identify and define the computing "
    "infrastructure and operations requirements appropriate to its solution. IT "
    "graduates should be able to work on large-scale computing systems."
)

# Programme Outcomes (PO) – hardcoded headings only; no custom PO.
# Used in multi-select "selected_pos" and in report generation.
PO_HEADINGS = [
    "Engineering Knowledge",
    "Problem Analysis",
    "Design / Development of Solutions",
    "Conduct investigations of complex problems",
    "Modern Tool Usage",
    "The Engineer and Society",
    "Environment and Sustainability",
    "Ethics",
    "Communication",
    "Project Management & Finance",
    "Lifelong Learning",
]


# ---------------- DATABASE ----------------

def init_db():
    conn = sqlite3.connect("database.db")
    c = conn.cursor()

    # Users table (with email and role)
    c.execute(
        """
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            username TEXT UNIQUE,
            email TEXT,
            password TEXT
        )
        """
    )

    # Ensure email / role columns exist on old databases
    c.execute("PRAGMA table_info(users)")
    user_cols = [row[1] for row in c.fetchall()]
    if "email" not in user_cols:
        c.execute("ALTER TABLE users ADD COLUMN email TEXT")
    if "role" not in user_cols:
        # Default everyone to 'Student' unless explicitly changed later
        c.execute("ALTER TABLE users ADD COLUMN role TEXT DEFAULT 'Student'")
        # Backfill any existing rows to have a role
        c.execute("UPDATE users SET role='Student' WHERE role IS NULL OR role=''")

    # Events table – stores all event metadata
    c.execute(
        """
        CREATE TABLE IF NOT EXISTS events (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            title TEXT,
            date TEXT,
            venue TEXT,
            department TEXT,
            description TEXT,
            event_photo TEXT,
            academic_year TEXT,
            resource_person TEXT,
            resource_designation TEXT,
            event_coordinator TEXT,
            event_time TEXT,
            event_type TEXT,
            permission_letter TEXT,
            invitation_letter TEXT,
            notice_letter TEXT,
            appreciation_letter TEXT,
            event_photos TEXT,
            attendance_photo TEXT,
            outcome_1 TEXT,
            outcome_2 TEXT,
            outcome_3 TEXT,
            feedback_data TEXT,
            pso1_selected INTEGER DEFAULT 0,
            pso2_selected INTEGER DEFAULT 0,
            selected_pos TEXT
        )
        """
    )

    # Add any missing columns for existing databases
    c.execute("PRAGMA table_info(events)")
    event_cols = [row[1] for row in c.fetchall()]
    new_event_columns = [
        ("event_photo", "TEXT"),
        ("academic_year", "TEXT"),
        ("resource_person", "TEXT"),
        ("resource_designation", "TEXT"),
        ("event_coordinator", "TEXT"),
        ("event_time", "TEXT"),
        ("event_type", "TEXT"),
        ("permission_letter", "TEXT"),
        ("invitation_letter", "TEXT"),
        ("notice_letter", "TEXT"),
        ("appreciation_letter", "TEXT"),
        ("event_photos", "TEXT"),
        ("attendance_photo", "TEXT"),
        ("outcome_1", "TEXT"),
        ("outcome_2", "TEXT"),
        ("outcome_3", "TEXT"),
        ("feedback_data", "TEXT"),
        ("pso1_selected", "INTEGER DEFAULT 0"),
        ("pso2_selected", "INTEGER DEFAULT 0"),
        ("selected_pos", "TEXT"),
    ]
    for col_name, col_type in new_event_columns:
        if col_name not in event_cols:
            c.execute(f"ALTER TABLE events ADD COLUMN {col_name} {col_type}")

    # Reports table – stores generated report info
    c.execute(
        """
        CREATE TABLE IF NOT EXISTS reports (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            event_id INTEGER,
            file_path TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
        """
    )

    # Ensure newer analytics-related columns exist on old databases
    c.execute("PRAGMA table_info(reports)")
    report_cols = [row[1] for row in c.fetchall()]
    if "status" not in report_cols:
        # Basic lifecycle tracking for reports; default everything to 'submitted'
        c.execute("ALTER TABLE reports ADD COLUMN status TEXT DEFAULT 'submitted'")
        c.execute("UPDATE reports SET status='submitted' WHERE status IS NULL OR status=''")

    conn.commit()
    conn.close()


def get_db():
    conn = sqlite3.connect("database.db")
    conn.row_factory = sqlite3.Row
    return conn


# --------------- FILE HELPERS ---------------

def _save_file(file_storage, subfolder, allow_pdf=False):
    """Save an uploaded file and return its relative path, or None."""
    if not file_storage or file_storage.filename == "":
        return None

    filename = secure_filename(file_storage.filename)
    ext = os.path.splitext(filename)[1].lower()
    image_exts = {".jpg", ".jpeg", ".png", ".gif"}
    allowed_exts = image_exts | ({".pdf"} if allow_pdf else set())
    if ext not in allowed_exts:
        return None

    os.makedirs(os.path.join(UPLOAD_FOLDER, subfolder), exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S_%f")
    filename = f"{timestamp}_{filename}"
    rel_path = os.path.join(UPLOAD_FOLDER, subfolder, filename)
    abs_path = os.path.join(rel_path)
    file_storage.save(abs_path)
    return rel_path.replace("\\", "/")


# ---------------- AUTH ----------------

def login_required(f):
    from functools import wraps
    @wraps(f)
    def wrap(*args, **kwargs):
        if "user_id" not in session:
            return redirect(url_for("login"))
        return f(*args, **kwargs)
    return wrap


def hod_required_api(f):
    """
    Simple decorator for HOD-only JSON APIs.
    Assumes user is already authenticated via @login_required.
    """
    from functools import wraps

    @wraps(f)
    def wrap(*args, **kwargs):
        if session.get("role") != "HOD":
            # Return a JSON 403 so frontend can handle gracefully
            return jsonify({"error": "Forbidden", "message": "HOD access required"}), 403
        return f(*args, **kwargs)

    return wrap


# ---------------- DOCX HELPERS ---------------- 
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def replace_placeholders(doc, replacements):
    """Replace placeholders in paragraphs and table cells."""
    def replace_in_paragraph(p):
        full_text = "".join(run.text for run in p.runs)
        modified = False
        for placeholder, value in replacements.items():
            if placeholder in full_text:
                # Clear all runs
                for run in p.runs:
                    run.text = ""
                # Add new text with replacement
                p.add_run(full_text.replace(placeholder, str(value or "")))
                modified = True
        return modified
    
    # Replace in paragraphs
    for p in doc.paragraphs:
        replace_in_paragraph(p)
    
    # Replace in table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph(p)
    
    # Replace in headers/footers
    for section in doc.sections:
        for header in [section.header, section.footer]:
            for p in header.paragraphs:
                replace_in_paragraph(p)
            for table in header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            replace_in_paragraph(p)

from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def insert_event_details_paragraph(doc, marker, event):
    for p in doc.paragraphs:
        if marker in p.text:
            # marker clear
            p.text = ""

            fields = [
                ("Academic Year", event["academic_year"]),
                ("Name of Event", event["title"]),
                ("Resource Person", event["resource_person"]),
                ("Event Type", event["event_type"]),
                ("Date", event["date"]),
                ("Time", event["event_time"]),
                ("Venue", event["venue"]),
                ("Department", event["department"]),
                ("Designation", event["resource_designation"]),
                ("Event Coordinator", event["event_coordinator"]),
            ]

            for label, value in reversed(fields):
                para = p.insert_paragraph_before()
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(2)
                para.paragraph_format.line_spacing = 1

                r1 = para.add_run(label)
                r1.bold = True
                r1.font.name = "Times New Roman"
                r1.font.size = Pt(12)

                para.add_run("\t")

                r2 = para.add_run(value or "")
                r2.bold = False
                r2.font.name = "Times New Roman"
                r2.font.size = Pt(12)

            return


def insert_full_page_image(doc, marker, path):
    if not path or not os.path.exists(path):
        return

    def process_paragraph(p):
        full_text = "".join(run.text for run in p.runs)
        if marker in full_text:
            for run in p.runs:
                run.text = ""
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(path, width=Inches(6))
            p.paragraph_format.keep_together = True
            return True
        return False

    
    for p in doc.paragraphs:
        if process_paragraph(p):
            return

    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if process_paragraph(p):
                        return

    
    for section in doc.sections:
        header = section.header
        for p in header.paragraphs:
            if process_paragraph(p):
                return

        for table in header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if process_paragraph(p):
                            return


from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def insert_event_photos(doc, marker, photos):
    if not photos:
        return

    for p in doc.paragraphs:
        if marker in p.text:
            p.text = p.text.replace(marker, "")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            for img in photos:
                if os.path.exists(img):
                    run = p.add_run()
                    run.add_picture(img, width=Inches(5))
                    p.add_run().add_break()

            return

from docx.enum.text import WD_BREAK
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def insert_attendance(doc, marker, path):
    if not path or not os.path.exists(path):
        return

    for p in doc.paragraphs:
        if marker in p.text:
            p.clear()

           
            run = p.add_run()
            run.add_break(WD_BREAK.PAGE)

            
            title_run = p.add_run("Attendance\n")
            title_run.bold = True
            title_run.font.size = Inches(0.25)  # visually bold/big title
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            
            p.paragraph_format.keep_with_next = True

            ext = os.path.splitext(path)[1].lower()

            
            if ext in [".jpg", ".jpeg", ".png"]:
                img_run = p.add_run()
                img_run.add_picture(path, width=Inches(8.0))
            else:
                p.add_run("\nAttendance attached as scanned document")

            return


from docx.shared import Pt

def set_cell_font(cell, bold=False):
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            run.bold = bold


def insert_feedback_table(doc, marker, feedback):
    if not feedback or len(feedback) < 1:
        return

    feedback = feedback[:10]

    for p in doc.paragraphs:
        if marker in p.text:
            p.clear()

            table = doc.add_table(rows=1, cols=3)
            try:
                table.style = "Table Grid"
            except KeyError:
                pass

            headers = ["Name", "Rating", "Comment"]
            for i, text in enumerate(headers):
                cell = table.rows[0].cells[i]
                cell.text = text
                set_cell_font(cell, bold=True)

            for fb in feedback:
                row = table.add_row().cells
                row[0].text = fb.get("name", "")
                row[1].text = fb.get("rating", "")
                row[2].text = fb.get("comment", "")

                for cell in row:
                    set_cell_font(cell)

            return


# ---------------- ROUTES ----------------

@app.route("/")
def home():
    return redirect(url_for("login"))


@app.route("/register", methods=["GET", "POST"])
def register():
    if request.method == "POST":
        username = request.form["username"]
        email = request.form.get("email", "")
        password = request.form["password"]
        role = request.form.get("role", "Student")  # Get selected role from form

        conn = get_db()
        existing = conn.execute(
            "SELECT id FROM users WHERE username=?", (username,)
        ).fetchone()
        if existing:
            conn.close()
            flash("Username already exists")
            return render_template("register.html")

        hashed = generate_password_hash(password)
        conn.execute(
            "INSERT INTO users (username, email, password, role) VALUES (?, ?, ?, ?)",
            (username, email, hashed, role),
        )
        conn.commit()
        conn.close()
        flash(f"Registration successful as {role}. Please login.")
        return redirect(url_for("login"))

    return render_template("register.html")


@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        username = request.form["username"]
        password = request.form["password"]
        selected_role = request.form.get("role", "Student")  # Get selected role from form

        conn = get_db()
        user = conn.execute(
            "SELECT * FROM users WHERE username=?", (username,)
        ).fetchone()
        conn.close()

        if user and check_password_hash(user["password"], password):
            # Get user's actual role from database
            user_role = user["role"] if "role" in user.keys() and user["role"] else "Student"
            
            # Verify that the selected role matches the user's role in database
            if selected_role != user_role:
                flash(f"Invalid role selection. Your account is registered as '{user_role}'. Please select the correct role.")
                return render_template("login.html")
            
            session["user_id"] = user["id"]
            session["username"] = user["username"]
            session["role"] = user_role
            return redirect(url_for("events"))
        flash("Invalid username or password")

    return render_template("login.html")


@app.route("/logout")
def logout():
    session.clear()
    return redirect(url_for("login"))


@app.route("/events")
@login_required
def events():
    conn = get_db()
    rows = conn.execute("SELECT * FROM events ORDER BY id DESC").fetchall()
    conn.close()

    events = []
    for e in rows:
        event = dict(e)   # convert sqlite row to dict
        # Parse event_photos JSON and ensure it's a list
        try:
            event["photos"] = json.loads(e["event_photos"]) if e["event_photos"] else []
        except (json.JSONDecodeError, TypeError):
            event["photos"] = []
        # Ensure photos is always a list
        if not isinstance(event["photos"], list):
            event["photos"] = []
        events.append(event)

    return render_template("events.html", events=events)



@app.route("/add_event", methods=["GET", "POST"])
@login_required
def add_event():
    # Check if this is an edit request (event_id in query params or form)
    event_id = request.args.get("event_id", type=int) or request.form.get("event_id", type=int)
    event = None
    
    # Load existing event data if editing
    if event_id:
        conn = get_db()
        event = conn.execute("SELECT * FROM events WHERE id=?", (event_id,)).fetchone()
        conn.close()
        if not event:
            flash("Event not found.")
            return redirect(url_for("events"))
        event = dict(event)  # Convert to dict for easier access
    
    if request.method == "POST":
        title = request.form["title"]
        date = request.form["date"]
        venue = request.form["venue"]
        department = request.form["department"]
        description = request.form.get("description", "")
        academic_year = request.form.get("academic_year", "")
        resource_person = request.form.get("resource_person", "")
        resource_designation = request.form.get("resource_designation", "")
        event_coordinator = request.form.get("event_coordinator", "")
        event_time = request.form.get("event_time", "")
        event_type = request.form.get("event_type", "")
        outcome_1 = request.form.get("outcome_1", "")
        outcome_2 = request.form.get("outcome_2", "")
        outcome_3 = request.form.get("outcome_3", "")

        # PSO selections from form (checkboxes). Store as integers 0/1 in DB.
        pso1_selected = 1 if request.form.get("pso1_selected") else 0
        pso2_selected = 1 if request.form.get("pso2_selected") else 0

        # Validate: At least one PSO must be selected.
        if not (pso1_selected or pso2_selected):
            flash("Please select at least one Programme Specific Outcome (PSO1 and/or PSO2).")
            # Redirect back to the same form (preserves event_id for edit flows)
            if event_id:
                return redirect(url_for("add_event", event_id=event_id))
            return redirect(url_for("add_event"))

        # Program Outcomes (PO): multi-select; only allow hardcoded headings (no custom).
        raw_pos = request.form.getlist("selected_pos")
        selected_pos = [p for p in raw_pos if p in PO_HEADINGS]
        if not selected_pos:
            flash("Please select at least one Program Outcome.")
            if event_id:
                return redirect(url_for("add_event", event_id=event_id))
            return redirect(url_for("add_event"))

        # Event photos (multiple) - preserve existing if no new uploads
        event_photos_paths = []
        if "event_photos" in request.files:
            for f in request.files.getlist("event_photos"):
                if f.filename:  # Only process if file was actually uploaded
                    saved = _save_file(f, "event_photos", allow_pdf=False)
                    if saved:
                        event_photos_paths.append(saved)
        
        # If editing and no new photos uploaded, keep existing photos
        if event_id and not event_photos_paths and event:
            try:
                existing_photos = json.loads(event["event_photos"]) if event.get("event_photos") else []
                if isinstance(existing_photos, list):
                    event_photos_paths = existing_photos
            except (json.JSONDecodeError, TypeError):
                pass

        event_photo_cover = event_photos_paths[0] if event_photos_paths else None

        # Scanned documents and attendance - preserve existing if no new uploads
        permission_letter = _save_file(
            request.files.get("permission_letter"), "attendance_photos", allow_pdf=True
        ) or (event["permission_letter"] if event_id and event else None)
        
        invitation_letter = _save_file(
            request.files.get("invitation_letter"), "attendance_photos", allow_pdf=True
        ) or (event["invitation_letter"] if event_id and event else None)
        
        notice_letter = _save_file(
            request.files.get("notice_letter"), "attendance_photos", allow_pdf=True
        ) or (event["notice_letter"] if event_id and event else None)
        
        appreciation_letter = _save_file(
            request.files.get("appreciation_letter"), "attendance_photos", allow_pdf=True
        ) or (event["appreciation_letter"] if event_id and event else None)
        
        attendance_photo = _save_file(
            request.files.get("attendance_photo"), "attendance_photos", allow_pdf=True
        ) or (event["attendance_photo"] if event_id and event else None)

        # Feedback
        feedback_names = request.form.getlist("feedback_name[]")
        feedback_ratings = request.form.getlist("feedback_rating[]")
        feedback_comments = request.form.getlist("feedback_comment[]")
        feedback_data = []
        for i in range(min(len(feedback_names), len(feedback_ratings), len(feedback_comments))):
            if (
                feedback_names[i].strip()
                or feedback_ratings[i].strip()
                or feedback_comments[i].strip()
            ):
                feedback_data.append(
                    {
                        "name": feedback_names[i],
                        "rating": feedback_ratings[i],
                        "comment": feedback_comments[i],
                    }
                )

        conn = get_db()
        
        # UPDATE if editing, INSERT if creating new
        if event_id:
            # Update existing event
            conn.execute(
                """
                UPDATE events SET
                    title=?, date=?, venue=?, department=?, description=?,
                    event_photo=?, academic_year=?, resource_person=?,
                    resource_designation=?, event_coordinator=?,
                    event_time=?, event_type=?,
                    permission_letter=?, invitation_letter=?, notice_letter=?,
                    appreciation_letter=?, event_photos=?, attendance_photo=?,
                    outcome_1=?, outcome_2=?, outcome_3=?, feedback_data=?,
                    pso1_selected=?, pso2_selected=?, selected_pos=?
                WHERE id=?
                """,
                (
                    title,
                    date,
                    venue,
                    department,
                    description,
                    event_photo_cover,
                    academic_year,
                    resource_person,
                    resource_designation,
                    event_coordinator,
                    event_time,
                    event_type,
                    permission_letter,
                    invitation_letter,
                    notice_letter,
                    appreciation_letter,
                    json.dumps(event_photos_paths) if event_photos_paths else None,
                    attendance_photo,
                    outcome_1,
                    outcome_2,
                    outcome_3,
                    json.dumps(feedback_data) if feedback_data else None,
                    pso1_selected,
                    pso2_selected,
                    json.dumps(selected_pos),
                    event_id,
                ),
            )
            flash("Event updated successfully.")
        else:
            # Insert new event
            conn.execute(
                """
                INSERT INTO events (
                    title, date, venue, department, description,
                    event_photo, academic_year, resource_person,
                    resource_designation, event_coordinator,
                    event_time, event_type,
                    permission_letter, invitation_letter, notice_letter,
                    appreciation_letter, event_photos, attendance_photo,
                    outcome_1, outcome_2, outcome_3, feedback_data,
                    pso1_selected, pso2_selected, selected_pos
                )
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    title,
                    date,
                    venue,
                    department,
                    description,
                    event_photo_cover,
                    academic_year,
                    resource_person,
                    resource_designation,
                    event_coordinator,
                    event_time,
                    event_type,
                    permission_letter,
                    invitation_letter,
                    notice_letter,
                    appreciation_letter,
                    json.dumps(event_photos_paths) if event_photos_paths else None,
                    attendance_photo,
                    outcome_1,
                    outcome_2,
                    outcome_3,
                    json.dumps(feedback_data) if feedback_data else None,
                    pso1_selected,
                    pso2_selected,
                    json.dumps(selected_pos),
                ),
            )
            flash("Event added successfully.")
        
        conn.commit()
        conn.close()
        return redirect(url_for("events"))

    # GET request - render form with existing data if editing
    # Parse feedback data if editing
    feedback_list = []
    if event and event.get("feedback_data"):
        try:
            feedback_list = json.loads(event["feedback_data"])
        except (json.JSONDecodeError, TypeError):
            feedback_list = []

    # Parse selected_pos for multi-select (edit mode)
    selected_pos_list = []
    if event and event.get("selected_pos"):
        try:
            selected_pos_list = json.loads(event["selected_pos"])
            if not isinstance(selected_pos_list, list):
                selected_pos_list = []
        except (json.JSONDecodeError, TypeError):
            selected_pos_list = []

    return render_template(
        "add_event.html",
        event=event,
        event_id=event_id,
        feedback_list=feedback_list,
        po_headings=PO_HEADINGS,
        selected_pos_list=selected_pos_list,
    )

@app.route("/event/<int:event_id>")
@login_required
def view_event(event_id):
    conn = get_db()
    event = conn.execute(
        "SELECT * FROM events WHERE id=?", (event_id,)
    ).fetchone()
    conn.close()

    if not event:
        flash("Event not found")
        return redirect(url_for("events"))

    photos = json.loads(event["event_photos"]) if event["event_photos"] else []
    feedback = json.loads(event["feedback_data"]) if event["feedback_data"] else []

    return render_template("view_event.html", event=dict(event), photos=photos, feedback=feedback)


@app.route("/delete_event/<int:event_id>", methods=["POST"])
@login_required
def delete_event(event_id):
    """Delete an event from the database."""
    conn = get_db()
    event = conn.execute("SELECT * FROM events WHERE id=?", (event_id,)).fetchone()
    
    if not event:
        conn.close()
        flash("Event not found.")
        return redirect(url_for("events"))
    
    # Delete the event
    conn.execute("DELETE FROM events WHERE id=?", (event_id,))
    conn.commit()
    conn.close()
    
    flash("Event deleted successfully.")
    return redirect(url_for("events"))

@app.route("/generate_report/<int:event_id>")
@login_required
def generate_report(event_id):

    conn = get_db()
    event = conn.execute("SELECT * FROM events WHERE id=?", (event_id,)).fetchone()
    conn.close()

    if not event:
        flash("Event not found")
        return redirect(url_for("events"))

    # Convert sqlite3.Row to plain dict so we can safely use .get()
    event = dict(event)

    doc = Document(TEMPLATE_PATH)

    # Build PSO section text based on selected PSOs for this event.
    # This will replace the {{PSO_SECTION}} placeholder in the template.
    selected_psos = []
    if event.get("pso1_selected"):
        selected_psos.append(PSO1_TEXT)
    if event.get("pso2_selected"):
        selected_psos.append(PSO2_TEXT)
    pso_section_text = "\n\n".join(selected_psos)

    # Build PO section: selected Program Outcomes as bullet list (for {{PO_SECTION}}).
    # Equivalent to: {% for po in selected_pos %} • {{po}} {% endfor %}
    selected_pos = []
    if event.get("selected_pos"):
        try:
            selected_pos = json.loads(event["selected_pos"])
            if not isinstance(selected_pos, list):
                selected_pos = []
        except (json.JSONDecodeError, TypeError):
            selected_pos = []
    po_section_text = "\n".join("• " + po for po in selected_pos)

    replace_placeholders(doc, {
        "{{academic_year}}": event["academic_year"],
        "{{date}}": event["date"],
        "{{event_name}}": event["title"],
        "{{event_type}}": event["event_type"],
        "{{event_date}}": event["date"],
        "{{event_time}}": event["event_time"],
        "{{venue}}": event["venue"],
        "{{department}}": event["department"],
        "{{resource_person}}": event["resource_person"],
        "{{resource_designation}}": event["resource_designation"],
        "{{event_coordinator}}": event["event_coordinator"],
        "{{event coordinator}}": event["event_coordinator"],
        # support both {{event_description}} and {{event description}} in template
        "{{event_description}}": event["description"],
        "{{event description}}": event["description"],
        "{{outcome_1}}": event["outcome_1"],
        "{{outcome_2}}": event["outcome_2"],
        "{{outcome_3}}": event["outcome_3"],
        "{{PSO_SECTION}}": pso_section_text,
        "{{PO_SECTION}}": po_section_text,
    })

    insert_event_details_paragraph(doc, "<<EVENT_DETAILS>>", event)

    insert_full_page_image(doc, "<<IMAGE_PERMISSION>>", event["permission_letter"])
    insert_full_page_image(doc, "<<IMAGE_INVITATION>>", event["invitation_letter"])
    insert_full_page_image(doc, "<<IMAGE_NOTICE>>", event["notice_letter"])
    insert_full_page_image(doc, "<<IMAGE_APPRECIATION>>", event["appreciation_letter"])

    photos = json.loads(event["event_photos"]) if event["event_photos"] else []
    insert_event_photos(doc, "<<EVENT_PHOTOS>>", photos)

    insert_attendance(doc, "<<ATTENDANCE_FILE>>", event["attendance_photo"])

    feedback = json.loads(event["feedback_data"]) if event["feedback_data"] else []
    insert_feedback_table(doc, "<<FEEDBACK_TABLE>>", feedback)

    filename = event["title"].replace(" ", "_") + ".docx"
    path = os.path.join(GENERATED_FOLDER, filename)
    doc.save(path)

    # Store report metadata
    conn = get_db()
    conn.execute(
        # status column has a default ('submitted'), but be explicit for clarity.
        "INSERT INTO reports (event_id, file_path, status) VALUES (?, ?, ?)",
        (event_id, path, "submitted"),
    )
    conn.commit()
    conn.close()

    return send_file(path, as_attachment=True)


@app.route("/reports")
@login_required
def reports():
    conn = get_db()
    rows = conn.execute(
        """
        SELECT r.id,
               r.file_path,
               r.created_at,
               e.title AS event_title,
               e.date  AS event_date
        FROM reports r
        JOIN events e ON r.event_id = e.id
        ORDER BY r.created_at DESC
        """
    ).fetchall()
    conn.close()
    # Shape data to match template expectations
    reports_data = []
    for r in rows:
        reports_data.append(
            {
                "id": r["id"],
                "file_path": r["file_path"],
                "created_at": r["created_at"],
                "event_title": r["event_title"],
                "event_date": r["event_date"],
                "created_by_name": "",  # no separate user table linkage here
            }
        )
    return render_template("reports.html", reports=reports_data)


@app.route("/download_report/<int:report_id>")
@login_required
def download_report(report_id):
    conn = get_db()
    row = conn.execute(
        "SELECT file_path FROM reports WHERE id=?", (report_id,)
    ).fetchone()
    conn.close()
    if not row or not row["file_path"] or not os.path.exists(row["file_path"]):
        flash("Report file not found.")
        return redirect(url_for("reports"))
    return send_file(row["file_path"], as_attachment=True)


@app.route("/api/hod/department-analysis")
@login_required
@hod_required_api
def hod_department_analysis_api():
    """
    Returns aggregated, department-wise report analytics for HOD users.

    Output format (example):
    {
        "departments": [
            {
                "department": "CSE",
                "total_reports": 10,
                "status_counts": {
                    "submitted": 5,
                    "approved": 3,
                    "rejected": 1,
                    "pending": 1
                }
            },
            ...
        ]
    }
    """
    conn = get_db()

    # Aggregate per department and per status. We keep the schema flexible while
    # still returning a shape that is easy for charting libraries to consume.
    rows = conn.execute(
        """
        SELECT
            e.department AS department,
            r.status     AS status,
            COUNT(r.id)  AS count
        FROM reports r
        JOIN events e ON r.event_id = e.id
        GROUP BY e.department, r.status
        ORDER BY e.department
        """
    ).fetchall()
    conn.close()

    # Build nested structure: department -> status_counts + total
    departments = {}
    for row in rows:
        dept = row["department"] or "Unknown"
        status = row["status"] or "unknown"
        count = row["count"] or 0

        if dept not in departments:
            departments[dept] = {
                "department": dept,
                "total_reports": 0,
                "status_counts": {},
            }
        departments[dept]["status_counts"][status] = (
            departments[dept]["status_counts"].get(status, 0) + count
        )
        departments[dept]["total_reports"] += count

    # Sorted list is a bit nicer for charts/UX
    department_list = sorted(
        departments.values(), key=lambda d: d["department"].lower()
    )

    return jsonify({"departments": department_list})


# ---------------- START ----------------

if __name__ == "__main__":
    os.makedirs(GENERATED_FOLDER, exist_ok=True)
    init_db()
    app.run(debug=True)
